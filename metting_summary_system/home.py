# -*- coding: utf-8 -*-
from collections import defaultdict

# -------------------------------------------------------------------------------
# Name:         home
# Description:  会议纪要工具
# Author:       shaver
# Date:         2025/7/1
# -------------------------------------------------------------------------------
import streamlit as st
from datetime import datetime, time
import base64
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 设置页面标题和布局
st.set_page_config(page_title="会议纪要系统", layout="wide")
# 引用SVG格式的图标
st.title(":coffee: 会议纪要系统")

# 添加侧边栏说明
with st.sidebar:
    st.header("使用说明")
    st.markdown("""
    1. 填写会议基本信息
    2. 添加讨论主题和事项
    3. 点击"生成WORD文档"
    4. 下载生成的WORD文档
    """)

# 将会议基本信息放入可折叠区域
with st.expander("会议基本信息", expanded=True):
    # 使用三列布局使表单更整齐
    col1, col2 = st.columns(2)

    with col1:
        # 选择会议日期
        meeting_date = st.date_input("会议日期", datetime.now().date())
        # 会议主持
        meeting_chair = st.text_input("会议主持人", "李超")

    with col2:
        # 会议记录人
        recorder = st.text_input("会议记录人", "")
        # 会议主题
        meeting_topic = st.text_input("会议主题", "日常工作讨论")

    # 会议地点单独一行保持对齐
    meeting_location = st.text_input("会议地点", "")

    # 参会人员放在最后，使用全宽度
    st.markdown("---")  # 添加分隔线
    options = ["张斌", "侯亚丽", "卢杰", "赵静", "李应龙", "肖涛", "任安安", "权昊", "李重瑛", "梁靖帆", "潘首文",
               "熊文江"]
    default_selected = options
    participants = st.multiselect("参会人员", options, default_selected)

# 初始化 session_state
if "topics" not in st.session_state:
    st.session_state.topics = {}  # 结构: { "主题1": [ {"task": "任务1", "person": "负责人1"}, ... ] }

# --- 添加新主题 ---
st.subheader("会议讨论内容")
new_topic = st.text_input("输入新主题名称:")
if st.button("添加主题") and new_topic:
    if new_topic not in st.session_state.topics:
        st.session_state.topics[new_topic] = []  # 初始化空记录列表
    else:
        st.warning("该主题已存在!")

# --- 显示所有主题及记录 ---
all_records = []
for topic_name, records in st.session_state.topics.items():
    with st.expander(f"主题: {topic_name}", expanded=True):
        # 主题标题行 - 更紧凑
        col1, col2 = st.columns([6, 1])
        with col1:
            st.markdown(f"**{topic_name}**")  # 使用markdown加粗代替subheader
        with col2:
            # 删除按钮更小
            if st.button("❌", key=f"del_topic_{topic_name}", help="删除整个主题"):
                del st.session_state.topics[topic_name]
                st.rerun()

        # --- 添加新记录表单 - 更紧凑 ---
        with st.form(key=f"add_record_{topic_name}"):
            cols = st.columns([4, 3, 1])  # 调整列比例
            with cols[0]:
                new_task = st.text_input("事项内容:", key=f"new_task_{topic_name}", label_visibility="collapsed",
                                         placeholder="事项内容")
            with cols[1]:
                new_person = st.text_input("负责人:", key=f"new_person_{topic_name}", label_visibility="collapsed",
                                           placeholder="负责人")
            with cols[2]:
                # 添加按钮更紧凑
                if st.form_submit_button("➕", help="添加记录"):
                    if new_task and new_person:
                        records.append({"task": new_task, "person": new_person, 'topic': topic_name})
                        st.rerun()
        # --- 显示所有记录 - 更紧凑 ---
        for i, record in enumerate(records):
            cols = st.columns([4, 3, 1])  # 调整列比例
            with cols[0]:
                # 修改事项 - 紧凑
                record["task"] = st.text_input(
                    "事项",
                    value=record["task"],
                    key=f"task_{topic_name}_{i}",
                    label_visibility="collapsed"
                )
            with cols[1]:
                # 修改负责人 - 紧凑
                record["person"] = st.text_input(
                    "负责人",
                    value=record["person"],
                    key=f"person_{topic_name}_{i}",
                    label_visibility="collapsed"
                )
            with cols[2]:
                # 删除按钮更小
                if st.button("❌", key=f"del_{topic_name}_{i}", help="删除此记录"):
                    records.pop(i)
                    st.rerun()
    all_records.extend(records)

# PDF生成函数
data = {}


def generate_word():
    # 处理all_records 按照topic分组
    grouped = defaultdict(list)
    for item in all_records:
        grouped[item["topic"]].append({"person": item["person"], "task": item["task"]})

    # 转换为目标格式
    sections = [{"topic": k, "topic_items": v} for k, v in grouped.items()]

    data = {
        "meeting_date": meeting_date,
        "meeting_location": meeting_location,
        "meeting_chair": meeting_chair,
        "participants": participants,
        "recorder": recorder,
        "meeting_topic": meeting_topic,
        "sections": sections
    }
    return data


# 生成PDF按钮
if st.button("生成WORD文档"):
    generate_word()
    if not st.session_state.topics:
        st.warning("请至少添加一个讨论主题")
    else:
        # 生成数据结构
        meeting_data = generate_word()

        # 创建Word文档
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '仿宋'
        font.size = Pt(10.5)

        # 添加标题
        title = doc.add_paragraph('会议纪要')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.style = doc.styles['Heading 1']

        # 添加基本信息
        doc.add_paragraph(f'会议时间: {meeting_data["meeting_date"]}')
        doc.add_paragraph(f'会议地点: {meeting_data["meeting_location"]}')
        doc.add_paragraph(f'会议主持: {meeting_data["meeting_chair"]}')

        # 添加参会人员
        participants_str = '、'.join(meeting_data["participants"])
        doc.add_paragraph(f'参会人员: {participants_str}')

        # 添加会议记录人员
        doc.add_paragraph(f'会议记录人员: {meeting_data["recorder"]}')

        # 添加空行
        doc.add_paragraph()

        # 添加会议议题
        topic = doc.add_paragraph(f'会议议题: {meeting_data["meeting_topic"]}')
        topic.style = doc.styles['Heading 2']

        # 添加空行
        doc.add_paragraph()

        # 添加各个部分
        for section in meeting_data["sections"]:
            # 添加主题
            section_topic = doc.add_paragraph(section["topic"])
            section_topic.style = doc.styles['Heading 3']

            # 添加主题子项
            if "topic_items" in section:
                for topic_item in section["topic_items"]:
                    # 添加内容
                    task = doc.add_paragraph(topic_item["task"], style='List Bullet')
                    # 添加责任人
                    if "person" in topic_item:
                        task.add_run(f'                  --{topic_item["person"]}')

        # 保存到临时文件
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        # 提供下载链接
        with open(tmp_path, 'rb') as f:
            file_bytes = f.read()
        os.unlink(tmp_path)  # 删除临时文件

        st.success("Word文档生成成功！")
        st.download_button(
            label="点击下载Word文档",
            data=file_bytes,
            file_name=f"会议纪要_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.malformations-office document.multiprocessing.document"
        )

# 添加页脚
st.markdown("---")
st.caption("© 2025 会议纪要系统 - 版本 1.0")
